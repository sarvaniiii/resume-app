import streamlit as st
import json
from datetime import datetime
from faker import Faker
from docx import Document
from docx.shared import Inches
import io

# Initialize Faker for sample data
fake = Faker()

class ResumeGenerator:
    def __init__(self):
        self.template = {
            "personal_info": {
                "name": "",
                "email": "",
                "phone": "",
                "address": "",
                "linkedin": "",
                "portfolio": ""
            },
            "summary": "",
            "experience": [],
            "education": [],
            "skills": [],
            "projects": []
        }
    
    def generate_sample_data(self):
        """Generate sample resume data"""
        self.template["personal_info"] = {
            "name": fake.name(),
            "email": fake.email(),
            "phone": fake.phone_number(),
            "address": fake.address().replace('\n', ', '),
            "linkedin": f"linkedin.com/in/{fake.user_name()}",
            "portfolio": f"{fake.user_name()}.portfolio.com"
        }
        
        self.template["summary"] = f"Experienced {fake.job()} with {fake.random_int(2, 10)} years in the industry. Skilled in various technologies and passionate about {fake.bs()}."
        
        # Generate sample experience
        self.template["experience"] = [
            {
                "company": fake.company(),
                "position": fake.job(),
                "duration": f"{fake.date_this_decade().strftime('%Y')} - Present",
                "description": fake.paragraph()
            }
            for _ in range(3)
        ]
        
        # Generate sample education
        self.template["education"] = [
            {
                "institution": fake.company() + " University",
                "degree": fake.random_element(["Bachelor of Science", "Master of Science", "PhD"]),
                "field": fake.random_element(["Computer Science", "Business Administration", "Engineering"]),
                "year": fake.year()
            }
            for _ in range(2)
        ]
        
        # Generate sample skills
        self.template["skills"] = [
            "Python", "JavaScript", "React", "Node.js", "SQL", "AWS",
            "Docker", "Git", "Machine Learning", "Data Analysis"
        ]
        
        # Generate sample projects
        self.template["projects"] = [
            {
                "name": fake.catch_phrase(),
                "description": fake.paragraph(),
                "technologies": [fake.random_element(self.template["skills"]) for _ in range(3)]
            }
            for _ in range(2)
        ]
        
        return self.template

class CoverLetterGenerator:
    def __init__(self):
        self.template = {
            "company_name": "",
            "hiring_manager": "",
            "position": "",
            "custom_content": ""
        }
    
    def generate_cover_letter(self, resume_data, cover_letter_data):
        """Generate cover letter based on resume and custom data"""
        
        cover_letter = f"""
{resume_data['personal_info']['name']}
{resume_data['personal_info']['address']}
{resume_data['personal_info']['phone']} | {resume_data['personal_info']['email']}

{datetime.now().strftime('%B %d, %Y')}

{cover_letter_data['hiring_manager']}
{cover_letter_data['company_name']}

Dear {cover_letter_data['hiring_manager'].split()[0] if cover_letter_data['hiring_manager'] else 'Hiring Manager'},

I am writing to express my enthusiastic interest in the {cover_letter_data['position']} position at {cover_letter_data['company_name']}. With my background in {', '.join(resume_data['skills'][:3])}, I am confident that I possess the skills and experience necessary to excel in this role.

{cover_letter_data['custom_content']}

My experience at {resume_data['experience'][0]['company'] if resume_data['experience'] else 'previous positions'} has equipped me with valuable skills in {', '.join(resume_data['skills'][:2])}, and I am excited about the opportunity to bring my expertise to {cover_letter_data['company_name']}.

Thank you for considering my application. I have attached my resume for your review and would welcome the opportunity to discuss how my skills and experiences align with your needs.

Sincerely,
{resume_data['personal_info']['name']}
"""
        return cover_letter

def create_word_document(content, doc_type="resume"):
    """Create a Word document from content"""
    doc = Document()
    
    if doc_type == "resume":
        # Add title
        title = doc.add_heading('Resume', 0)
        
        # Add personal info
        personal_info = content['personal_info']
        doc.add_paragraph(f"Name: {personal_info['name']}")
        doc.add_paragraph(f"Email: {personal_info['email']}")
        doc.add_paragraph(f"Phone: {personal_info['phone']}")
        doc.add_paragraph(f"Address: {personal_info['address']}")
        if personal_info['linkedin']:
            doc.add_paragraph(f"LinkedIn: {personal_info['linkedin']}")
        if personal_info['portfolio']:
            doc.add_paragraph(f"Portfolio: {personal_info['portfolio']}")
        
        # Add summary
        doc.add_heading('Professional Summary', level=1)
        doc.add_paragraph(content['summary'])
        
        # Add experience
        doc.add_heading('Experience', level=1)
        for exp in content['experience']:
            exp_heading = doc.add_heading(f"{exp['position']} at {exp['company']}", level=2)
            doc.add_paragraph(f"Duration: {exp['duration']}")
            doc.add_paragraph(exp['description'])
        
        # Add education
        doc.add_heading('Education', level=1)
        for edu in content['education']:
            edu_heading = doc.add_heading(f"{edu['degree']} in {edu['field']}", level=2)
            doc.add_paragraph(f"Institution: {edu['institution']}")
            doc.add_paragraph(f"Year: {edu['year']}")
        
        # Add skills
        doc.add_heading('Skills', level=1)
        skills_para = doc.add_paragraph()
        for skill in content['skills']:
            skills_para.add_run(f"• {skill} ")
    
    else:  # cover letter
        doc.add_paragraph(content)
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.set_page_config(
        page_title="Resume & Cover Letter Generator",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("📄 Resume & Cover Letter Generator")
    st.markdown("Create professional resumes and cover letters in minutes!")
    
    # Initialize generators
    resume_gen = ResumeGenerator()
    cover_letter_gen = CoverLetterGenerator()
    
    # Sidebar for navigation
    st.sidebar.title("Navigation")
    app_mode = st.sidebar.selectbox("Choose what to generate:", 
                                   ["Resume Builder", "Cover Letter Generator"])
    
    if app_mode == "Resume Builder":
        st.header("Resume Builder")
        
        # Sample data button
        if st.button("Generate Sample Data"):
            sample_data = resume_gen.generate_sample_data()
            st.session_state.resume_data = sample_data
            st.success("Sample data generated! Fill the form below or edit as needed.")
        
        # Initialize session state
        if 'resume_data' not in st.session_state:
            st.session_state.resume_data = resume_gen.template
        
        # Personal Information
        st.subheader("Personal Information")
        col1, col2 = st.columns(2)
        
        with col1:
            st.session_state.resume_data['personal_info']['name'] = st.text_input(
                "Full Name", 
                st.session_state.resume_data['personal_info']['name']
            )
            st.session_state.resume_data['personal_info']['email'] = st.text_input(
                "Email", 
                st.session_state.resume_data['personal_info']['email']
            )
            st.session_state.resume_data['personal_info']['phone'] = st.text_input(
                "Phone", 
                st.session_state.resume_data['personal_info']['phone']
            )
        
        with col2:
            st.session_state.resume_data['personal_info']['address'] = st.text_input(
                "Address", 
                st.session_state.resume_data['personal_info']['address']
            )
            st.session_state.resume_data['personal_info']['linkedin'] = st.text_input(
                "LinkedIn", 
                st.session_state.resume_data['personal_info']['linkedin']
            )
            st.session_state.resume_data['personal_info']['portfolio'] = st.text_input(
                "Portfolio", 
                st.session_state.resume_data['personal_info']['portfolio']
            )
        
        # Professional Summary
        st.subheader("Professional Summary")
        st.session_state.resume_data['summary'] = st.text_area(
            "Summary",
            st.session_state.resume_data['summary'],
            height=100
        )
        
        # Experience
        st.subheader("Work Experience")
        for i, exp in enumerate(st.session_state.resume_data['experience']):
            st.markdown(f"**Experience {i+1}**")
            col1, col2 = st.columns(2)
            
            with col1:
                exp['company'] = st.text_input(f"Company {i+1}", exp['company'], key=f"company_{i}")
                exp['position'] = st.text_input(f"Position {i+1}", exp['position'], key=f"position_{i}")
            
            with col2:
                exp['duration'] = st.text_input(f"Duration {i+1}", exp['duration'], key=f"duration_{i}")
            
            exp['description'] = st.text_area(f"Description {i+1}", exp['description'], key=f"desc_{i}")
        
        # Add/Remove experience entries
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Add Experience"):
                st.session_state.resume_data['experience'].append({
                    "company": "", "position": "", "duration": "", "description": ""
                })
                st.rerun()
        
        with col2:
            if len(st.session_state.resume_data['experience']) > 0:
                if st.button("Remove Last Experience"):
                    st.session_state.resume_data['experience'].pop()
                    st.rerun()
        
        # Skills
        st.subheader("Skills")
        skills_text = st.text_area(
            "Enter skills (one per line or comma-separated)",
            "\n".join(st.session_state.resume_data['skills'])
        )
        if skills_text:
            # Handle both newline and comma separation
            skills_list = [skill.strip() for skill in skills_text.replace(',', '\n').split('\n') if skill.strip()]
            st.session_state.resume_data['skills'] = skills_list
        
        # Education
        st.subheader("Education")
        for i, edu in enumerate(st.session_state.resume_data['education']):
            st.markdown(f"**Education {i+1}**")
            col1, col2 = st.columns(2)
            
            with col1:
                edu['institution'] = st.text_input(f"Institution {i+1}", edu['institution'], key=f"inst_{i}")
                edu['degree'] = st.text_input(f"Degree {i+1}", edu['degree'], key=f"degree_{i}")
            
            with col2:
                edu['field'] = st.text_input(f"Field {i+1}", edu['field'], key=f"field_{i}")
                edu['year'] = st.text_input(f"Year {i+1}", edu['year'], key=f"year_{i}")
        
        # Add/Remove education entries
        col1, col2 = st.columns(2)
        with col1:
            if st.button("Add Education"):
                st.session_state.resume_data['education'].append({
                    "institution": "", "degree": "", "field": "", "year": ""
                })
                st.rerun()
        
        with col2:
            if len(st.session_state.resume_data['education']) > 0:
                if st.button("Remove Last Education"):
                    st.session_state.resume_data['education'].pop()
                    st.rerun()
        
        # Preview and Download
        st.subheader("Preview & Download")
        
        if st.button("Generate Resume Preview"):
            # Display preview
            st.json(st.session_state.resume_data)
            
            # Create download button for Word document
            doc_buffer = create_word_document(st.session_state.resume_data, "resume")
            st.download_button(
                label="📥 Download Resume (Word)",
                data=doc_buffer,
                file_name=f"resume_{st.session_state.resume_data['personal_info']['name'].replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
    else:  # Cover Letter Generator
        st.header("Cover Letter Generator")
        
        # Check if resume data exists
        if 'resume_data' not in st.session_state:
            st.warning("Please fill out your resume information first to generate a cover letter.")
            if st.button("Generate Sample Resume Data"):
                st.session_state.resume_data = resume_gen.generate_sample_data()
                st.rerun()
            return
        
        # Cover letter form
        st.subheader("Cover Letter Details")
        
        col1, col2 = st.columns(2)
        
        with col1:
            company_name = st.text_input("Company Name")
            position = st.text_input("Position Applying For")
        
        with col2:
            hiring_manager = st.text_input("Hiring Manager Name", placeholder="Leave blank for 'Hiring Manager'")
            custom_content = st.text_area(
                "Why you're interested in this position?",
                height=150,
                placeholder="Describe why you're a good fit for this role and company..."
            )
        
        cover_letter_data = {
            "company_name": company_name,
            "hiring_manager": hiring_manager or "Hiring Manager",
            "position": position,
            "custom_content": custom_content
        }
        
        if st.button("Generate Cover Letter"):
            if company_name and position:
                cover_letter = cover_letter_gen.generate_cover_letter(
                    st.session_state.resume_data, 
                    cover_letter_data
                )
                
                st.subheader("Cover Letter Preview")
                st.text_area("Preview", cover_letter, height=400)
                
                # Create download button for Word document
                doc_buffer = create_word_document(cover_letter, "cover_letter")
                st.download_button(
                    label="📥 Download Cover Letter (Word)",
                    data=doc_buffer,
                    file_name=f"cover_letter_{company_name.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Text download
                st.download_button(
                    label="📥 Download Cover Letter (Text)",
                    data=cover_letter,
                    file_name=f"cover_letter_{company_name.replace(' ', '_')}.txt",
                    mime="text/plain"
                )
            else:
                st.error("Please fill in Company Name and Position fields.")

if __name__ == "__main__":
    main()